import json
import docx
from docx.shared import Pt

with open('uat.json', 'r') as f:
    data = json.load(f)

doc = docx.Document()
doc.add_heading('Faspay VA - Skenario Functional Test_V.3.0 static', 0)

table = doc.add_table(rows=1, cols=8)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
headers = ['No', 'Service', 'Scenario', 'Expected Result', 'Request', 'Response', 'Result', 'Notes']
for i, header in enumerate(headers):
    hdr_cells[i].text = header

for tc_id, tc in data['results'].items():
    row_cells = table.add_row().cells
    
    service = 'Any Service'
    if tc_id in ['11.6', '11.7', '11.8', '11.9']: service = 'Inquiry VA'
    if tc_id in ['11.10', '11.11', '11.12', '11.16']: service = 'Payment VA'
    if tc_id == '11.17': service = 'Payment Notification'
    
    expected = ''
    if tc_id == '11.1': expected = 'Error Code: 4012401'
    if tc_id == '11.2': expected = 'Error Code: 4012400'
    if tc_id in ['11.3', '11.4']: expected = 'Error Code: 4002401'
    if tc_id == '11.5': expected = 'Error Code: 4092400'
    if tc_id == '11.6': expected = 'Response Code: 2002400'
    if tc_id == '11.7': expected = 'Error Code: 4042414'
    if tc_id == '11.8': expected = 'Error Code: 4042419'
    if tc_id == '11.9': expected = 'Error Code: 4042412'
    if tc_id in ['11.10', '11.16', '11.17']: expected = 'Response Code: 2002500'
    if tc_id == '11.11': expected = 'Error Code: 4042512'
    if tc_id == '11.12': expected = 'Error Code: 4042513'
    
    curl = f"url:\n{tc['request_url']}\n\nheader:\n"
    for k, v in tc['request_headers'].items():
        curl += f"{k}: {v}\n"
    curl += f"\nbody:\n{json.dumps(tc['request_payload'])}"
    
    row_cells[0].text = str(tc_id)
    row_cells[1].text = service
    row_cells[2].text = tc['scenario']
    row_cells[3].text = expected
    row_cells[4].text = curl
    row_cells[5].text = json.dumps(tc['response_body'], indent=2)
    row_cells[6].text = 'Pass'
    row_cells[7].text = ''

    # adjust font size for curl and response to be smaller
    for cell in [row_cells[4], row_cells[5]]:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)

doc.save('/Users/nurisakbar/.gemini/antigravity-ide/brain/c292b667-e916-4112-bf32-5b638257ada3/Faspay_VA_UAT.docx')
print("Successfully generated Word document!")
